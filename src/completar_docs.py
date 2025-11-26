
from docx import Document
import json
import os

# Archivo base del contrato
PLANTILLA = "CONTRATO-MATRICULA-2025-La-Florida.docx"
JSON_DATOS = "datos_contratos.json"

def completar_contrato(datos, plantilla):
    contract_src = './assets/docs'
    contract_template = os.path.join(contract_src, plantilla)
    doc = Document(contract_template)

    # Reemplazar texto en párrafos
    for p in doc.paragraphs:
        if "Don(ña):" in p.text:
            p.text = f"Don(ña): {datos['apoderado']['nombre']} en su calidad de: {datos['apoderado']['calidad']}"
        if "de profesión u oficio:" in p.text:
            p.text = (f"de profesión u oficio: {datos['apoderado']['profesion']} RUT: {datos['apoderado']['rut']} "
                      f"domiciliado(a) en (calle): {datos['apoderado']['domicilio']['calle']} N°: {datos['apoderado']['domicilio']['numero']} "
                      f"Casa: {datos['apoderado']['domicilio']['casa']} Depto: {datos['apoderado']['domicilio']['depto']} Comuna: {datos['apoderado']['domicilio']['comuna']}")
        if "Nombre Completo" in p.text:
            p.text = ""  # Limpiar texto original

    # Insertar alumnos
    for alumno in datos['alumnos']:
        doc.add_paragraph(f"Nombre Completo: {alumno['nombre']}      Curso: {alumno['curso']}")

    # Insertar firma
    doc.add_paragraph(f"Nombre: {datos['firma']['nombre']}")
    doc.add_paragraph(f"C.I.: {datos['firma']['ci']}")
    doc.add_paragraph(f"La Florida {datos['firma']['fecha']}")

    return doc

def generar_contratos():
    # Leer datos desde el archivo JSON
    data_src = './assets/json'
    json_data = os.path.join(data_src, JSON_DATOS)
    with open(json_data, "r", encoding="utf-8") as f:
        contratos = json.load(f)["contratos"]

    # Crear carpeta de salida
    os.makedirs("Contratos_Completados", exist_ok=True)

    # Generar contratos completados
    #for i, contrato in enumerate(contratos, start=1):
    for contrac_name in contratos:
        doc_completado = completar_contrato(contratos[contrac_name], contrac_name)
        output_name = contrac_name[:-5] + "_completado.docx"
        #output_file = os.path.join("Contratos_Completados", f"CONTRATO-MATRICULA-COMPLETADO-{i}.docx")
        output_file = os.path.join("Contratos_Completados", output_name)
        doc_completado.save(output_file)
        print(f"✅ Contrato completado guardado en: {output_file}")

if __name__ == "__main__":
    generar_contratos()
